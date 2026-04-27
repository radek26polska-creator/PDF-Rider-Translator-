"""
PDF to DOCX Converter - Idealne odwzorowanie layoutu
Zachowuje: tekst (pogrubienie, kursywa, rozmiary), tabele, pozycjonowanie, obrazy
Dodano: OCR dla skanowanych PDF, lepsze wykrywanie tabel i formatowania
"""

import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import pdfplumber
import io
import os
import tempfile
from typing import List, Optional, Dict, Any
import logging

# OCR (opcjonalne)
try:
    import pytesseract
    from PIL import Image
    import numpy as np
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

logger = logging.getLogger(__name__)


class PDFtoDOCXConverter:
    def __init__(self, pdf_path: str, docx_path: str, use_ocr: bool = False, lang: str = "pol+eng"):
        self.pdf_path = pdf_path
        self.docx_path = docx_path
        self.doc = Document()
        self.pdf_doc = None
        self.use_ocr = use_ocr and OCR_AVAILABLE
        self.ocr_lang = lang
        self.page_width = 0
        self.page_height = 0

        # Ustaw marginesy (1 cal)
        sections = self.doc.sections
        for section in sections:
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)

    def open_pdf(self):
        self.pdf_doc = fitz.open(self.pdf_path)
        first_page = self.pdf_doc[0]
        self.page_width = first_page.rect.width
        self.page_height = first_page.rect.height

    def close(self):
        if self.pdf_doc:
            self.pdf_doc.close()

    def convert(self) -> bool:
        self.open_pdf()

        for page_num in range(len(self.pdf_doc)):
            page = self.pdf_doc[page_num]
            self._convert_page(page)

        self.doc.save(self.docx_path)
        self.close()
        return True

    def _convert_page(self, page):
        """Konwertuj stronę z uwzględnieniem OCR jeśli potrzebne"""
        # Sprawdź, czy strona ma tekst
        text_dict = page.get_text("dict")
        text_blocks = [b for b in text_dict.get("blocks", []) if b["type"] == 0 and b.get("lines")]

        # Jeśli brak tekstu i OCR włączony – użyj OCR
        if not text_blocks and self.use_ocr:
            self._convert_page_with_ocr(page)
            return

        # Normalna konwersja z tekstem
        image_list = page.get_images(full=True)
        tables = self._extract_tables_with_pdfplumber(page.number)

        elements = []

        # Tekst – grupowanie bloków
        for block in text_blocks:
            elements.append({"type": "text", "y": block["bbox"][1], "data": block})

        # Tabele – rozszerzone wykrywanie
        tables_data = self._extract_tables_enhanced(page)
        for table in tables_data:
            elements.append({"type": "table", "y": table["bbox"][1], "data": table})

        # Obrazy
        for img_info in image_list:
            xref = img_info[0]
            rect = self._get_image_rect(page, xref)
            if rect:
                elements.append({"type": "image", "y": rect[1], "data": {"xref": xref, "rect": rect}})

        # Sortuj po Y
        elements.sort(key=lambda e: e["y"])

        for element in elements:
            if element["type"] == "text":
                self._add_text_block(element["data"])
            elif element["type"] == "table":
                self._add_table(element["data"]["table_data"])
            elif element["type"] == "image":
                self._add_image_from_page(page, element["data"]["xref"])

        if page.number < len(self.pdf_doc) - 1:
            self.doc.add_page_break()

    def _convert_page_with_ocr(self, page):
        """Konwertuj stronę używając OCR (dla skanów)"""
        logger.info(f"OCR strona {page.number + 1}")
        
        # Renderuj stronę jako obraz
        mat = fitz.Matrix(2, 2)  # wyższa rozdzielczość dla OCR
        pix = page.get_pixmap(matrix=mat)
        img_data = pix.tobytes("png")
        
        # OCR przez pytesseract
        image = Image.open(io.BytesIO(img_data))
        ocr_text = pytesseract.image_to_pdf_or_hocr(image, extension='pdf', lang=self.ocr_lang)
        
        # Zapisz tymczasowy PDF z OCR
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp:
            tmp.write(ocr_text)
            tmp_path = tmp.name
        
        # Otwórz OCR PDF i wyodrębnij tekst
        ocr_doc = fitz.open(tmp_path)
        ocr_page = ocr_doc[0]
        text = ocr_page.get_text()
        ocr_doc.close()
        os.unlink(tmp_path)
        
        if text.strip():
            self.doc.add_paragraph(text)
        else:
            self.doc.add_paragraph("[Strona nie zawiera tekstu]")

    def _extract_tables_enhanced(self, page) -> List[dict]:
        """Rozszerzona ekstrakcja tabel – łączy pdfplumber z heurystykami"""
        tables = []
        try:
            with pdfplumber.open(self.pdf_path) as pdf:
                pdf_page = pdf.pages[page.number]
                
                # Wykryj tabele
                table_objs = pdf_page.find_tables()
                for table_obj in table_objs:
                    table_data = table_obj.extract()
                    if table_data and len(table_data) > 0:
                        # Oczyść dane
                        cleaned = []
                        for row in table_data:
                            cleaned_row = [str(cell).strip() if cell else "" for cell in row]
                            cleaned.append(cleaned_row)
                        
                        tables.append({
                            "data": cleaned,
                            "bbox": table_obj.bbox,
                            "page": page.number
                        })
        except Exception as e:
            logger.warning(f"Błąd pdfplumber strona {page.number}: {e}")
        
        return tables

    def _add_text_block(self, block):
        lines = block.get("lines", [])
        for line in lines:
            paragraph = self.doc.add_paragraph()

            for span in line.get("spans", []):
                text = span["text"].strip()
                if not text:
                    continue

                run = paragraph.add_run(text)

                # Czcionka
                font = span.get("font", "").lower()
                if "bold" in font or "bold" in span.get("font", ""):
                    run.bold = True
                if "italic" in font or "italic" in span.get("font", ""):
                    run.italic = True

                # Rozmiar
                size = span.get("size", 12)
                run.font.size = Pt(size)

                # Kolor
                color = span.get("color", (0, 0, 0))
                if color and len(color) >= 3:
                    r, g, b = [int(c * 255) for c in color[:3]]
                    run.font.color.rgb = RGBColor(r, g, b)

            # Interlinia
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    def _add_table(self, table_data: List[List[str]]):
        if not table_data or len(table_data) == 0:
            return

        rows = len(table_data)
        cols = max(len(row) for row in table_data)

        if rows == 0 or cols == 0:
            return

        table = self.doc.add_table(rows=rows, cols=cols)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.style = 'Table Grid'

        for r_idx, row in enumerate(table_data):
            for c_idx, cell_text in enumerate(row):
                cell = table.cell(r_idx, c_idx)
                cell.text = str(cell_text) if cell_text else ""
                # Nagłówek
                if r_idx == 0:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.bold = True

    def _add_image_from_page(self, page, xref: int):
        try:
            base = self.pdf_doc.extract_image(xref)
            image_stream = io.BytesIO(base["image"])
            # Zachowaj proporcje
            self.doc.add_picture(image_stream, width=Inches(2.0))
        except Exception as e:
            print(f"Błąd obrazu: {e}")

    def _get_image_rect(self, page, xref: int) -> Optional[tuple]:
        try:
            rects = page.get_image_rects(xref)
            return rects[0] if rects else None
        except Exception:
            return None


def convert_pdf_to_docx(pdf_path: str, docx_path: str, use_ocr: bool = False, ocr_lang: str = "pol+eng") -> bool:
    try:
        converter = PDFtoDOCXConverter(pdf_path, docx_path, use_ocr=use_ocr, lang=ocr_lang)
        converter.convert()
        return True
    except Exception as e:
        print(f"Błąd: {e}")
        return False
